import socket
import concurrent.futures
import nmap
import tkinter as tk
from tkinter import scrolledtext, ttk, filedialog, messagebox, StringVar
import requests
import subprocess
import random
import time
import threading
from scapy.all import IP, ICMP, sr1, conf, ARP, Ether, srp
import folium
import webbrowser
from geopy.geocoders import Nominatim
import os
import tkintermapview
import ttkbootstrap as ttkb
from ttkbootstrap.constants import *
from ttkbootstrap import Window, Style
import customtkinter as ctk
from tkinter import IntVar
from tkinter.scrolledtext import ScrolledText
from fpdf import FPDF
from docx import Document
import phonenumbers
from phonenumbers import geocoder, carrier, timezone
import socket 
from tkintermapview import TkinterMapView
import csv
from customtkinter import CTk, CTkButton, CTkEntry, CTkLabel, CTkFrame, CTkTabview
from email_validator import validate_email, EmailNotValidError
import whois
import dns.resolver

# Global variables to store results
open_ports = {}
closed_ports = {}
nm = nmap.PortScanner()  # Create an instance of the nmap.PortScanner
start_time = None  # Variable to store the start time of the scan

# Function to get geolocation for an IP
def get_geolocation(ip):
    try:
        response = requests.get(f"https://ipinfo.io/{ip}/json?token=your api")
        if response.status_code == 200:
            data = response.json()
            return {
                "IP": ip,
                "City": data.get("city", "N/A"),
                "Region": data.get("region", "N/A"),
                "Country": data.get("country", "N/A"),
                "Latitude, Longitude": data.get("loc", "N/A"),
                "ISP": data.get("org", "N/A")
            }
        else:
            return {"Error": "Could not fetch geolocation data."}
    except requests.exceptions.RequestException as e:
        return {"Error": str(e)}

# Function to update the output text area in a thread-safe manner
def update_output(text_area, message):
    text_area.insert(tk.END, message)
    text_area.see(tk.END)

# Function to scan a single port
def scan_port(ip, port, scan_text_area, progress_var, start_port, end_port, stop_event):
    try:
        if stop_event.is_set():  # Stop immediately if the event is set
            update_output(scan_text_area, f"Port scanning stopped at port {port}\n")
            return
        family = socket.AF_INET6 if ':' in ip else socket.AF_INET
        with socket.socket(family, socket.SOCK_STREAM) as sock:
            sock.settimeout(1)
            result = sock.connect_ex((ip, port))
            if result == 0:
                service_name = socket.getservbyport(port, 'tcp')
                open_ports[port] = service_name  # Store open ports with service names
                update_output(scan_text_area, f"Port {port} is OPEN (Service: {service_name})\n")
            else:
                closed_ports[port] = "CLOSED"  # Store closed ports
                update_output(scan_text_area, f"Port {port} is CLOSED\n")

        # Evasion technique: random delay to avoid detection
        delay = random.uniform(0.1, 0.5)
        time.sleep(delay)

        progress = int(((port - start_port + 1) / (end_port - start_port + 1)) * 100)
        progress_var.set(progress)
        
    except Exception as err:
        update_output(scan_text_area, f"Could not connect to {ip} on port {port}: {err}\n")

# Function to scan ports in the provided range
def scan_ports(ip, start_port, end_port, scan_text_area, stop_event, progress_var):
    global open_ports, closed_ports, start_time
    open_ports = {}
    closed_ports = {}

    geolocation_info = get_geolocation(ip)
    display_geolocation(ip, geolocation_info)

    update_output(scan_text_area, f"\nScanning IP: {ip}\n")
    update_output(scan_text_area, f"Scanning ports from {start_port} to {end_port}...\n")

    start_time = time.time()  # Start the timer

    with concurrent.futures.ThreadPoolExecutor(max_workers=50) as executor:
        futures = []
        for port in range(start_port, end_port + 1):
            if stop_event.is_set():
                update_output(scan_text_area, "Scan cancelled.\n")
                break
            futures.append(executor.submit(scan_port, ip, port, scan_text_area, progress_var, start_port, end_port, stop_event))
        
        # Wait for all futures to complete
        for future in concurrent.futures.as_completed(futures):
            if stop_event.is_set():
                update_output(scan_text_area,  "Scan cancelled.\n")
                break

    # Sort and display open ports
    if open_ports:
        update_output(scan_text_area, "\nOpen Ports (Sorted):\n")
        for port in sorted(open_ports.keys()):
            service = open_ports[port]
            update_output(scan_text_area, f"Port {port} (Service: {service})\n")

    elapsed_time = time.time() - start_time  # Calculate elapsed time
    update_output(scan_text_area, f"\nScan Complete! Elapsed time: {elapsed_time:.2f} seconds\n")
    progress_var.set(100)
    

    # Identify the device type after scanning
    identify_device(ip)

    # Perform vulnerability scan on open ports
    if open_ports:
        update_output(vuln_text_area, "\nOpen Ports Found:\n")
        for port, service in open_ports.items():
            update_output(vuln_text_area, f"Port: {port} (Service: {service})\n")
        perform_vulnerability_scan(ip, vuln_text_area)

# Function to perform a real-world traceroute
def real_traceroute(ip, traceroute_text_area):
    update_output(traceroute_text_area, f"\nTraceroute to {ip}:\n")
    try:
        # Using Scapy for traceroute
        max_hops = 30
        for ttl in range(1, max_hops + 1):
            pkt = IP(dst=ip, ttl=ttl) / ICMP()
            reply = sr1(pkt, verbose=0, timeout=1)
            if reply is None:
                update_output(traceroute_text_area, f"{ttl}: * * * Request timed out.\n")
                continue
            elif reply.type == 0:  # Echo reply
                update_output(traceroute_text_area, f"{ttl}: {reply.src} (Reached destination)\n")
                break
            else:
                update_output(traceroute_text_area, f"{ttl}: {reply.src}\n")
    except Exception as e:
        update_output(traceroute_text_area, f"Traceroute command failed: {e}\n")

# Function to identify the device type using Nmap
def identify_device(ip):
    try:
        nm.scan(ip, arguments='-O')  # Run Nmap with OS detection
        display_device_info(ip)
    except Exception as e:
        update_output(traceroute_text_area, f"OS Detection failed: {e}\n")

def display_device_info(ip):
    update_output(traceroute_text_area, f"\nDevice Identification Results for {ip}:\n")
    device_info_found = False
    device_types = {
        "mobile": ["Android", "iPhone", "iOS", "Windows Phone", "Mobile", "Smartphone"],
        "tablet": ["iPad", "Android Tablet", "Tablet"],
        "server": ["Linux", "Windows Server", "FreeBSD", "Unix", "Apache", "Nginx", "Tomcat"],
        "desktop": ["Windows", "Mac OS X", "Linux", "Ubuntu", "Fedora"],
        "iot": ["Raspberry Pi", "IoT", "Smart"],
        "printer": ["Printer", "HP", "Canon"],
        "camera": ["IP Camera", "Webcam"],
        "network_device": ["Router", "Switch", "Access Point"],
        "website": ["HTTP", "HTTPS", "Web", "Web Server"]
    }

    if 'osclass' in nm[ip]:
        for os in nm[ip]['osclass']:
            update_output(traceroute_text_area, f"Device Type: {os['osfamily']} - {os['osgen']}\n")
            device_info_found = True
            # Check if it's a mobile, server, or other device
            identify_and_display_device_type(os['osfamily'], device_types)

    elif 'osmatch' in nm[ip]:
        for os in nm[ip]['osmatch']:
            update_output(traceroute_text_area, f"Device Type: {os['name']}\n")
            device_info_found = True
            # Check if it's a mobile, server, or other device
            identify_and_display_device_type(os['name'], device_types)

    if not device_info_found:
        update_output(traceroute_text_area, "No device type information found.\n")

def identify_and_display_device_type(os_info, device_types):
    for device_type, keywords in device_types.items():
        if any(keyword in os_info for keyword in keywords):
            update_output(traceroute_text_area, f"Device identified as: {device_type.capitalize()}\n")
            break  # Stop checking once the device type is identified

# Function to manually identify device type when button is pressed
def identify_device_button():
    ip = ip_entry.get()
    identify_device(ip)


# Function to perform a vulnerability scan on open ports
def perform_vulnerability_scan(ip, vuln_text_area):
    update_output(vuln_text_area, "\nPerforming deep vulnerability scan on open ports...\n")
    for port, service in open_ports.items():
        try:
            nm.scan(ip, arguments=f"-p {port} --script=vuln")  # Scan with vulnerability scripts
            vulnerability_info = nm[ip]['tcp'][port]['script']
            update_output(vuln_text_area, f"Vulnerabilities for Port {port} ({service}):\n")
            for script, output in vulnerability_info.items():
                update_output(vuln_text_area, f"{script}: {output}\n")
        except Exception as e:
            update_output(vuln_text_area, f"Error scanning port {port} for vulnerabilities: {e}\n")

# Function to display geolocation information in the GUI
def display_geolocation(ip, geolocation_info):
    geo_text_area.delete('1.0', tk.END)
    geo_text_area.insert(tk.END, "Geolocation Information:\n")
    for key, value in geolocation_info.items():
        geo_text_area.insert(tk.END, f"{key}: {value}\n")

# Function to clear all text areas and inputs
def clear_all():
    ip_entry.delete(0, tk.END)
    start_port_entry.delete(0, tk.END)
    end_port_entry.delete(0, tk.END)
    geo_text_area.delete('1.0', tk.END)
    scan_text_area.delete('1.0', tk.END)
    vuln_text_area.delete('1.0', tk.END)
    traceroute_text_area.delete('1.0', tk.END)
    progress_var.set(0)

# Function to start the scan
def start_scan():
    global stop_event
    stop_event = threading.Event()
    ip = ip_entry.get()
    start_port = int(start_port_entry.get())
    end_port = int(end_port_entry.get())
    threading.Thread(target=scan_ports, args=(ip, start_port, end_port, scan_text_area, stop_event, progress_var)).start()

# Function to stop the scan
def stop_scan():
    stop_event.set()  # Stop the ongoing scan
    update_output(scan_text_area, "Scan stopped.\n")

# Function to start the traceroute
def start_traceroute():
    ip = ip_entry.get()
    threading.Thread(target=real_traceroute, args=(ip, traceroute_text_area)).start()

def show_map():
    location = map_entry.get()
    if location:
        map_widget.set_address(location)  # Center the map to the entered location
    else:
        map_widget.set_position(20, 0)  # Default location (latitude and longitude)

# Placeholder for dynamically gathered information
def gathered_info():
    """
    Gathers information from the text areas of the GUI.
    Returns a dictionary with the gathered information.
    """
    geo_info = geo_text_area.get("1.0", "end-1c")
    scan_results = scan_text_area.get("1.0", "end-1c")
    vuln_results = vuln_text_area.get("1.0", "end-1c")
    traceroute_results = traceroute_text_area.get("1.0", "end-1c")
    
    info = {
        "Geolocation Information": geo_info,
        "Port Scanning Results": scan_results,
        "Vulnerability Scanning Results": vuln_results,
        "Traceroute Results": traceroute_results,
    }
    
    # Combine into a single string for saving
    report = "\n\n".join(f"{section}:\n{content}" for section, content in info.items())
    return report

# Save as TXT
def save_as_txt():
    file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")])
    if file_path:
        with open(file_path, "w") as file:
            file.write(gathered_info())
        messagebox.showinfo("Success", "Report saved as TXT.")

# Save as PDF
def save_as_pdf():
    file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf")])
    if file_path:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for line in gathered_info().split("\n"):
            pdf.cell(200, 10, txt=line, ln=True)
        pdf.output(file_path)
        messagebox.showinfo("Success", "Report saved as PDF.")

# Save as DOCX
def save_as_docx():
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    if file_path:
        doc = Document()
        doc.add_heading("Network Scanning Report", level=1)
        for line in gathered_info().split("\n"):
            if line.strip():  # Avoid adding blank lines
                doc.add_paragraph(line)
        doc.save(file_path)
        messagebox.showinfo("Success", "Report saved as DOCX.")

# Dropdown selection handler
def handle_dropdown(choice):
    if choice == "Save as TXT":
        save_as_txt()
    elif choice == "Save as PDF":
        save_as_pdf()
    elif choice == "Save as DOCX":
        save_as_docx()

# Fuction to get mac from ip and display
def get_mac(ip):
    arp_request = ARP(pdst=ip)
    broadcast = Ether(dst="ff:ff:ff:ff:ff:ff")
    arp_request_broadcast = broadcast/arp_request
    
    # Send the request and capture the response
    answered_list = srp(arp_request_broadcast, timeout=1, verbose=False)[0]
    
    if answered_list:
        return answered_list[0][1].hwsrc
    else:
        return None

def display_mac():
    """ Function to get MAC and display it in the result label """
    ip = ip_entry.get()  # Get the IP address from the entry field
    mac_address = get_mac(ip)
    
    if mac_address:
        result_label.configure(text=f"MAC Address: {mac_address}")
    else:
        result_label.configure(text="MAC Address not found!")
        
# Function to query MACVendors API and fetch vendor info
def get_mac_vendor(mac_address):
    url = f"https://api.macvendors.com/{mac_address}"
    
    try:
        # Send GET request to the MACVendors API
        response = requests.get(url)
        
        # If the request is successful
        if response.status_code == 200:
            return response.text  # Vendor name is returned as a plain text response
        else:
            return f"Error: {response.status_code} - {response.text}"

    except requests.exceptions.RequestException as e:
        return f"Request failed: {e}"

# Function to update the output text area with MAC address vendor info
def get_mac_info():
    mac_address = mac_entry.get()
    if mac_address:
        vendor_info = get_mac_vendor(mac_address)
        output_text.configure(state="normal")
        output_text.delete(1.0, "end")  # Clear previous output
        output_text.insert("1.0", f"MAC Address: {mac_address}\nVendor: {vendor_info}\n")
        output_text.configure(state="disabled")  # Disable to prevent editing
    else:
        output_text.configure(state="normal")
        output_text.delete(1.0, "end")
        output_text.insert("1.0", "Please enter a valid MAC address.\n")
        output_text.configure(state="disabled")

# List of country codes
country_codes = {
    "Afghanistan (+93)": "+93",
    "Albania (+355)": "+355",
    "Algeria (+213)": "+213",
    "Andorra (+376)": "+376",
    "Angola (+244)": "+244",
    "Antarctica (+672)": "+672",
    "Argentina (+54)": "+54",
    "Armenia (+374)": "+374",
    "Australia (+61)": "+61",
    "Austria (+43)": "+43",
    "Azerbaijan (+994)": "+994",
    "Bahamas (+1 242)": "+1 242",
    "Bahrain (+973)": "+973",
    "Bangladesh (+880)": "+880",
    "Barbados (+1 246)": "+1 246",
    "Belarus (+375)": "+375",
    "Belgium (+32)": "+32",
    "Belize (+501)": "+501",
    "Benin (+229)": "+229",
    "Bhutan (+975)": "+975",
    "Bolivia (+591)": "+591",
    "Bosnia and Herzegovina (+387)": "+387",
    "Botswana (+267)": "+267",
    "Brazil (+55)": "+55",
    "Brunei (+673)": "+673",
    "Bulgaria (+359)": "+359",
    "Burkina Faso (+226)": "+226",
    "Burundi (+257)": "+257",
    "Cambodia (+855)": "+855",
    "Cameroon (+237)": "+237",
    "Canada (+1)": "+1",
    "Chad (+235)": "+235",
    "Chile (+56)": "+56",
    "China (+86)": "+86",
    "Colombia (+57)": "+57",
    "Comoros (+269)": "+269",
    "Congo (+242)": "+242",
    "Costa Rica (+506)": "+506",
    "Croatia (+385)": "+385",
    "Cuba (+53)": "+53",
    "Cyprus (+357)": "+357",
    "Czech Republic (+420)": "+420",
    "Denmark (+45)": "+45",
    "Djibouti (+253)": "+253",
    "Dominica (+1 767)": "+1 767",
    "Dominican Republic (+1 809)": "+1 809",
    "Ecuador (+593)": "+593",
    "Egypt (+20)": "+20",
    "El Salvador (+503)": "+503",
    "Estonia (+372)": "+372",
    "Ethiopia (+251)": "+251",
    "Fiji (+679)": "+679",
    "Finland (+358)": "+358",
    "France (+33)": "+33",
    "Gabon (+241)": "+241",
    "Gambia (+220)": "+220",
    "Georgia (+995)": "+995",
    "Germany (+49)": "+49",
    "Ghana (+233)": "+233",
    "Greece (+30)": "+30",
    "Grenada (+1 473)": "+1 473",
    "Guatemala (+502)": "+502",
    "Guinea (+224)": "+224",
    "Guyana (+592)": "+592",
    "Haiti (+509)": "+509",
    "Honduras (+504)": "+504",
    "Hungary (+36)": "+36",
    "Iceland (+354)": "+354",
    "India (+91)": "+91",
    "Indonesia (+62)": "+62",
    "Iran (+98)": "+98",
    "Iraq (+964)": "+964",
    "Ireland (+353)": "+353",
    "Israel (+972)": "+972",
    "Italy (+39)": "+39",
    "Jamaica (+1 876)": "+1 876",
    "Japan (+81)": "+81",
    "Jordan (+962)": "+962",
    "Kazakhstan (+7)": "+7",
    "Kenya (+254)": "+254",
    "Kiribati (+686)": "+686",
    "Kuwait (+965)": "+965",
    "Kyrgyzstan (+996)": "+996",
    "Laos (+856)": "+856",
    "Latvia (+371)": "+371",
    "Lebanon (+961)": "+961",
    "Lesotho (+266)": "+266",
    "Liberia (+231)": "+231",
    "Libya (+218)": "+218",
    "Liechtenstein (+423)": "+423",
    "Lithuania (+370)": "+370",
    "Luxembourg (+352)": "+352",
    "Madagascar (+261)": "+261",
    "Malawi (+265)": "+265",
    "Malaysia (+60)": "+60",
    "Maldives (+960)": "+960",
    "Mali (+223)": "+223",
    "Malta (+356)": "+356",
    "Marshall Islands (+692)": "+692",
    "Mauritania (+222)": "+222",
    "Mauritius (+230)": "+230",
    "Mexico (+52)": "+52",
    "Micronesia (+691)": "+691",
    "Moldova (+373)": "+373",
    "Monaco (+377)": "+377",
    "Mongolia (+976)": "+976",
    "Montenegro (+382)": "+382",
    "Morocco (+212)": "+212",
    "Mozambique (+258)": "+258",
    "Myanmar (+95)": "+95",
    "Namibia (+264)": "+264",
    "Nauru (+674)": "+674",
    "Nepal (+977)": "+977",
    "Netherlands (+31)": "+31",
    "New Zealand (+64)": "+64",
    "Nicaragua (+505)": "+505",
    "Niger (+227)": "+227",
    "Nigeria (+234)": "+234",
    "Norway (+47)": "+47",
    "Oman (+968)": "+968",
    "Pakistan (+92)": "+92",
    "Palau (+680)": "+680",
    "Panama (+507)": "+507",
    "Papua New Guinea (+675)": "+675",
    "Paraguay (+595)": "+595",
    "Peru (+51)": "+51",
    "Philippines (+63)": "+63",
    "Poland (+48)": "+48",
    "Portugal (+351)": "+351",
    "Qatar (+974)": "+974",
    "Romania (+40)": "+40",
    "Russia (+7)": "+7",
    "Rwanda (+250)": "+250",
    "Samoa (+685)": "+685",
    "San Marino (+378)": "+378",
    "Saudi Arabia (+966)": "+966",
    "Senegal (+221)": "+221",
    "Serbia (+381)": "+381",
    "Seychelles (+248)": "+248",
    "Singapore (+65)": "+65",
    "Slovakia (+421)": "+421",
    "Slovenia (+386)": "+386",
    "Somalia (+252)": "+252",
    "South Africa (+27)": "+27",
    "South Korea (+82)": "+82",
    "Spain (+34)": "+34",
    "Sri Lanka (+94)": "+94",
    "Sudan (+249)": "+249",
    "Sweden (+46)": "+46",
    "Switzerland (+41)": "+41",
    "Syria (+963)": "+963",
    "Taiwan (+886)": "+886",
    "Tajikistan (+992)": "+992",
    "Tanzania (+255)": "+255",
    "Thailand (+66)": "+66",
    "Togo (+228)": "+228",
    "Tonga (+676)": "+676",
    "Trinidad and Tobago (+1 868)": "+1 868",
    "Tunisia (+216)": "+216",
    "Turkey (+90)": "+90",
    "Uganda (+256)": "+256",
    "Ukraine (+380)": "+380",
    "United Arab Emirates (+971)": "+971",
    "United Kingdom (+44)": "+44",
    "United States (+1)": "+1",
    "Uruguay (+598)": "+598",
    "Uzbekistan (+998)": "+998",
    "Vanuatu (+678)": "+678",
    "Vatican City (+379)": "+379",
    "Venezuela (+58)": "+58",
    "Vietnam (+84)": "+84",
    "Yemen (+967)": "+967",
    "Zambia (+260)": "+260",
    "Zimbabwe (+263)": "+263"
}

# Function to fetch geolocation data using NumVerify API
def fetch_geolocation(phone_number):
    API_KEY = "your api"  # Replace with your NumVerify API key
    API_URL = f"http://apilayer.net/api/validate?access_key={API_KEY}&number={phone_number}"
    
    try:
        response = requests.get(API_URL)
        data = response.json()
        
        if data.get("valid"):
            location = data.get("location", "Unknown")
            line_type = data.get("line_type", "Unknown")
            carrier_name = data.get("carrier", "Unknown")
            return f"Location: {location}\nLine Type: {line_type}\nCarrier: {carrier_name}"
        else:
            return "Number not valid or no additional data available."
    except Exception as e:
        return f"Error fetching geolocation: {e}"

# Function to lookup phone number details
def lookup_phone():
    selected_code = country_code_var.get()
    user_number = entry.get()
    phone_number = selected_code + user_number.lstrip("+")

    if not user_number:
        messagebox.showerror("Error", "Please enter the phone number.")
        return
    
    try:
        parsed_number = phonenumbers.parse(phone_number)
        
        if not phonenumbers.is_valid_number(parsed_number):
            messagebox.showerror("Invalid", "The phone number is not valid.")
            return
        
        country = geocoder.description_for_number(parsed_number, "en")
        service_provider = carrier.name_for_number(parsed_number, "en")
        time_zones = timezone.time_zones_for_number(parsed_number)
        geo_data = fetch_geolocation(phone_number)
        
        output_box.configure(state="normal")
        output_box.delete("1.0", tk.END)
        output_box.insert(tk.END, f"Country: {country}\n", "bold")
        output_box.insert(tk.END, f"Carrier: {service_provider}\n", "info")
        output_box.insert(tk.END, f"Time Zone(s): {', '.join(time_zones)}\n", "info")
        output_box.insert(tk.END, f"{geo_data}\n", "geo")
        output_box.configure(state="disabled")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

# Function to clear all fields
def clear_fields():
    entry.delete(0, tk.END)
    output_box.configure(state="normal")
    output_box.delete("1.0", tk.END)
    output_box.configure(state="disabled")

# Function for scanning connected IPs
def reverse_dns_lookup(ip):
    try:
        host, _, _ = socket.gethostbyaddr(ip)  # Try to resolve IP to a hostname
        return host
    except socket.herror:
        return None  # If no hostname found, return None

def scan_network():
    ip_range = ip_range_entry.get()  # Get IP range from GUI input
    nm = nmap.PortScanner()
    
    try:
        nm.scan(hosts=ip_range, arguments="-sn")  # Perform a ping scan
        active_ips = [host for host in nm.all_hosts() if nm[host].state() == "up"]
        
        # Display active IPs and hostnames in the GUI
        active_ips_text.delete(1.0, ctk.END)  # Clear previous results
        for ip in active_ips:
            hostname = reverse_dns_lookup(ip)  # Perform reverse DNS lookup

            active_ips_text.insert(ctk.END, f"IP: {ip}\n")
            active_ips_text.insert(ctk.END, f"Hostname: {hostname if hostname else 'N/A'}\n\n")
    except Exception as e:
        messagebox.showerror("Error", f"Scan failed: {e}")

def lscan_network():
    ip_range = lip_range_entry.get()  # Get IP range from GUI input
    nm = nmap.PortScanner()
    
    try:
        nm.scan(hosts=ip_range, arguments="-sn")  # Perform a ping scan
        active_ips = [host for host in nm.all_hosts() if nm[host].state() == "up"]
        
        # Display active IPs and hostnames in the GUI
        lactive_ips_text.delete(1.0, ctk.END)  # Clear previous results
        for ip in active_ips:
            hostname = reverse_dns_lookup(ip)  # Perform reverse DNS lookup

            lactive_ips_text.insert(ctk.END, f"IP: {ip}\n")
            lactive_ips_text.insert(ctk.END, f"Hostname: {hostname if hostname else 'N/A'}\n\n")
    except Exception as e:
        messagebox.showerror("Error", f"Scan failed: {e}")

# Global Variables
flight_lines = []
countries = [
    "Afghanistan", "Albania", "Algeria", "Andorra", "Angola", "Antigua and Barbuda",
    "Argentina", "Armenia", "Australia", "Austria", "Azerbaijan", "Bahamas", "Bahrain",
    "Bangladesh", "Barbados", "Belarus", "Belgium", "Belize", "Benin", "Bhutan", "Bolivia",
    "Bosnia and Herzegovina", "Botswana", "Brazil", "Brunei", "Bulgaria", "Burkina Faso",
    "Burundi", "Cabo Verde", "Cambodia", "Cameroon", "Canada", "Central African Republic",
    "Chad", "Chile", "China", "Colombia", "Comoros", "Congo (Congo-Brazzaville)",
    "Costa Rica", "Croatia", "Cuba", "Cyprus", "Czech Republic (Czechia)", "Democratic Republic of the Congo",
    "Denmark", "Djibouti", "Dominica", "Dominican Republic", "Ecuador", "Egypt", "El Salvador",
    "Equatorial Guinea", "Eritrea", "Estonia", "Eswatini (fmr. 'Swaziland')", "Ethiopia", "Fiji", "Finland",
    "France", "Gabon", "Gambia", "Georgia", "Germany", "Ghana", "Greece", "Grenada", "Guatemala", "Guinea",
    "Guinea-Bissau", "Guyana", "Haiti", "Honduras", "Hungary", "Iceland", "India", "Indonesia", "Iran", "Iraq",
    "Ireland", "Israel", "Italy", "Ivory Coast", "Jamaica", "Japan", "Jordan", "Kazakhstan", "Kenya",
    "Kiribati", "Korea, North", "Korea, South", "Kuwait", "Kyrgyzstan", "Laos", "Latvia", "Lebanon",
    "Lesotho", "Liberia", "Libya", "Liechtenstein", "Lithuania", "Luxembourg", "Madagascar", "Malawi",
    "Malaysia", "Maldives", "Mali", "Malta", "Marshall Islands", "Mauritania", "Mauritius", "Mexico", "Micronesia",
    "Moldova", "Monaco", "Mongolia", "Montenegro", "Morocco", "Mozambique", "Myanmar (formerly Burma)",
    "Namibia", "Nauru", "Nepal", "Netherlands", "New Zealand", "Nicaragua", "Niger", "Nigeria", "North Macedonia",
    "Norway", "Oman", "Pakistan", "Palau", "Panama", "Papua New Guinea", "Paraguay", "Peru", "Philippines",
    "Poland", "Portugal", "Qatar", "Romania", "Russia", "Rwanda", "Saint Kitts and Nevis", "Saint Lucia",
    "Saint Vincent and the Grenadines", "Samoa", "San Marino", "Sao Tome and Principe", "Saudi Arabia",
    "Senegal", "Serbia", "Seychelles", "Sierra Leone", "Singapore", "Slovakia", "Slovenia", "Solomon Islands",
    "Somalia", "South Africa", "South Sudan", "Spain", "Sri Lanka", "Sudan", "Suriname", "Sweden", "Switzerland",
    "Syria", "Taiwan", "Tajikistan", "Tanzania", "Thailand", "Timor-Leste", "Togo", "Tonga", "Trinidad and Tobago",
    "Tunisia", "Turkey", "Turkmenistan", "Tuvalu", "Uganda", "Ukraine", "United Arab Emirates", "United Kingdom",
    "United States of America", "Uruguay", "Uzbekistan", "Vanuatu", "Vatican City", "Venezuela", "Vietnam",
    "Yemen", "Zambia", "Zimbabwe"
]


# Function to fetch live flight data
def f_fetch_flight_data():
    url = "https://opensky-network.org/api/states/all"
    try:
        response = requests.get(url)
        response.raise_for_status()
        data = response.json()
        flights = data.get("states", [])
        
        # Log all callsigns for debugging
        all_callsigns = [flight[1] for flight in flights if flight[1]]
        
        return flights
    except requests.exceptions.RequestException as e:
        f_log_message(f"Error fetching flight data: {e}")
        messagebox.showerror("Error", "Failed to fetch flight data.")
        return []


# Function to filter flights by country
def f_filter_flights_by_country(flights, country):
    filtered = []
    skipped_flights = []
    for flight in flights:
        if flight[2] == country:
            start_latitude, start_longitude = flight[6], flight[5]
            end_latitude, end_longitude = flight[9], flight[8]

            if None in [start_latitude, start_longitude, end_latitude, end_longitude]:
                skipped_flights.append((flight[1], "Missing coordinates"))
                continue

            if not (-90 <= start_latitude <= 90 and -180 <= start_longitude <= 180):
                skipped_flights.append((flight[1], "Start coordinates out of range"))
                continue
            if not (-90 <= end_latitude <= 90 and -180 <= end_longitude <= 180):
                skipped_flights.append((flight[1], "End coordinates out of range"))
                continue

            filtered.append(flight)

    if skipped_flights:
        for flight, reason in skipped_flights:
            f_log_message(f"Skipping flight {flight} due to: {reason}")

    return filtered


# Function to log messages to the GUI
def f_log_message(message):
    f_log_text.configure(state="normal")
    f_log_text.insert("end", message + "\n")
    f_log_text.see("end")
    f_log_text.configure(state="disabled")


# Function to update map
def f_update_map_live():
    global flight_lines
    selected_country = f_country_menu.get()
    if selected_country == "Select a Country":
        f_log_message("No country selected.")
        messagebox.showwarning("Warning", "Please select a country.")
        return

    f_loading_label.configure(text="Loading flights...")
    f_view_button.configure(state="disabled")

    def fetch_and_update():
        try:
            flights = f_fetch_flight_data()
            filtered_flights = f_filter_flights_by_country(flights, selected_country)

            f_map_widget.delete_all_marker()
            for line in flight_lines:
                line.delete()
            flight_lines.clear()

            if not filtered_flights:
                f_log_message(f"No flights found for {selected_country}.")
                messagebox.showinfo("No Flights", f"No flights found for {selected_country}.")
                return

            for flight in filtered_flights:
                callsign = flight[1] or "Unknown"
                start_latitude, start_longitude = flight[6], flight[5]
                altitude = flight[7] or "N/A"
                speed = flight[9] or "N/A"

                # Set marker with flight details as tooltip
                marker = f_map_widget.set_marker(
                    start_latitude, start_longitude,
                    text=f"{callsign}\nAltitude: {altitude} m\nSpeed: {speed} m/s"
                )
                


            f_log_message(f"Displayed {len(filtered_flights)} flights for {selected_country}.")
            messagebox.showinfo("Flights Displayed", f"Displayed {len(filtered_flights)} flights for {selected_country}.")
        except Exception as e:
            f_log_message(f"Error updating map: {e}")
            messagebox.showerror("Error", "Failed to update the map.")
        finally:
            f_loading_label.configure(text="")
            f_view_button.configure(state="normal")

    threading.Thread(target=fetch_and_update, daemon=True).start()

def f_search_callsign(callsign):
    if not callsign:
        f_log_message("No callsign entered.")
        messagebox.showwarning("Warning", "Please enter a callsign.")
        return

    f_loading_label.configure(text="Searching for flight...")
    f_search_button.configure(state="disabled")

    def fetch_and_display():
        try:
            flights = f_fetch_flight_data()
            matching_flight = next((flight for flight in flights if callsign in (flight[1] or "")), None)

            if not matching_flight:
                f_log_message(f"No flight found with callsign '{callsign}'.")
                messagebox.showinfo("No Flight", f"No flight found with callsign '{callsign}'.")
                return

            start_latitude, start_longitude = matching_flight[6], matching_flight[5]
            altitude = matching_flight[7] or "N/A"
            speed = matching_flight[9] or "N/A"

            f_map_widget.delete_all_marker()
            f_map_widget.set_marker(
                start_latitude, start_longitude,
                text=f"Callsign: {callsign}\nAltitude: {altitude} m\nSpeed: {speed} m/s"
            )
            f_map_widget.set_position(start_latitude, start_longitude)
            f_map_widget.set_zoom(6)

            f_log_message(f"Displayed flight '{callsign}' at ({start_latitude}, {start_longitude}) with altitude {altitude} m and speed {speed} m/s.")
            messagebox.showinfo("Flight Found", f"Displayed flight '{callsign}' on the map.")
        except Exception as e:
            f_log_message(f"Error searching flight: {e}")
            messagebox.showerror("Error", "Failed to search for the flight.")
        finally:
            f_loading_label.configure(text="")
            f_search_button.configure(state="normal")

    threading.Thread(target=fetch_and_display, daemon=True).start()

def f_clear_all():
    # Clear country selection
    f_country_menu.set("Select a Country")
    
    # Clear callsign entry
    callsign_var.set("")
    
    # Clear log text
    f_log_text.configure(state="normal")
    f_log_text.delete("1.0", "end")
    f_log_text.configure(state="disabled")
    
    # Clear map markers
    f_map_widget.delete_all_marker()
    
    # Log the action
    f_log_message("All inputs and map cleared.")

# Function to handle button click
def f_view_flights():
    f_update_map_live()

# Enable typing and auto-complete by binding an event
def f_on_country_input(event):
    value = country_var.get().lower()
    filtered_countries = [country for country in countries if value in country.lower()]
    country_menu['values'] = filtered_countries if filtered_countries else countries

# Global Variables
API_KEY = "your api"
WEATHER_API_KEY = "your api"

coordinates_list = []
markers = []
location_names = []

# Functions
def m_search_location():
    location = m_search_entry.get()
    if location:
        url = f"https://geocode.search.hereapi.com/v1/geocode?q={location}&apiKey={API_KEY}"
        response = requests.get(url).json()

        if response.get("items"):
            position = response["items"][0]["position"]
            lat, lon = position["lat"], position["lng"]
            m_map_widget.set_position(lat, lon)
            m_map_widget.set_zoom(10)
            m_result_text.delete(1.0, tk.END)
            m_result_text.insert(tk.END, f"Coordinates: {lat}, {lon}\n")
            coordinates_list.append((lat, lon))
        else:
            m_result_text.delete(1.0, tk.END)
            m_result_text.insert(tk.END, "Location not found.")

def m_get_weather():
    location = m_search_entry.get()
    if not location:
        m_result_text.delete(1.0, tk.END)
        m_result_text.insert(tk.END, "Enter a location to fetch weather.\n")
        return

    url = f"http://api.weatherapi.com/v1/current.json?key={WEATHER_API_KEY}&q={location}&aqi=no"

    try:
        response = requests.get(url).json()
        if "error" in response:
            m_result_text.delete(1.0, tk.END)
            m_result_text.insert(tk.END, f"Error: {response['error']['message']}\n")
        else:
            weather = response["current"]
            condition = weather["condition"]["text"]
            temp_c = weather["temp_c"]
            feelslike_c = weather["feelslike_c"]
            humidity = weather["humidity"]
            wind_kph = weather["wind_kph"]

            m_result_text.delete(1.0, tk.END)
            m_result_text.insert(tk.END, f"Weather in {location}:\n")
            m_result_text.insert(tk.END, f"  Condition: {condition}\n")
            m_result_text.insert(tk.END, f"  Temperature: {temp_c}°C (Feels like: {feelslike_c}°C)\n")
            m_result_text.insert(tk.END, f"  Humidity: {humidity}%\n")
            m_result_text.insert(tk.END, f"  Wind Speed: {wind_kph} kph\n")
    except Exception as e:
        m_result_text.delete(1.0, tk.END)
        m_result_text.insert(tk.END, f"Failed to fetch weather data: {str(e)}\n")

def m_reverse_geocode():
    center = m_map_widget.get_position()
    lat, lon = center[0], center[1]

    url = f"https://revgeocode.search.hereapi.com/v1/revgeocode?at={lat},{lon}&apiKey={API_KEY}"
    response = requests.get(url).json()

    if response.get("items"):
        address = response["items"][0]["address"]["label"]
        m_result_text.delete(1.0, tk.END)
        m_result_text.insert(tk.END, f"Address: {address}\n")
    else:
        m_result_text.delete(1.0, tk.END)
        m_result_text.insert(tk.END, "Address not found.")

def m_add_marker():
    if not m_search_entry.get():
        tk.messagebox.showerror("Error", "Search location cannot be empty!")
        return

    location_name = m_search_entry.get()
    geocode_url = f"https://geocode.search.hereapi.com/v1/geocode?q={location_name}&apiKey={API_KEY}"
    response = requests.get(geocode_url).json()

    if response.get("items"):
        position = response["items"][0]["position"]
        lat, lng = position["lat"], position["lng"]
        coordinates_list.append((lat, lng))
        location_names.append(location_name)
        m_map_widget.set_marker(lat, lng, text=location_name)
        m_fit_bounds_manual()  # Automatically fit bounds after adding a marker
    else:
        tk.messagebox.showerror("Error", "Location not found!")

    location_name = m_search_entry.get()
    geocode_url = f"https://geocode.search.hereapi.com/v1/geocode?q={location_name}&apiKey={API_KEY}"
    response = requests.get(geocode_url).json()

    if response.get("items"):
        position = response["items"][0]["position"]
        lat, lng = position["lat"], position["lng"]
        coordinates_list.append((lat, lng))
        location_names.append(location_name)
        m_map_widget.set_marker(lat, lng, text=location_name)
    else:
        tk.messagebox.showerror("Error", "Location not found!")

def m_calculate_route():
    if len(coordinates_list) < 2:
        m_result_text.delete(1.0, tk.END)
        m_result_text.insert(tk.END, "Add at least two locations to calculate the route.\n")
        return

    total_distance = 0
    total_duration = 0
    results = []

    for i in range(len(coordinates_list) - 1):
        start_coords = f"{coordinates_list[i][0]},{coordinates_list[i][1]}"
        end_coords = f"{coordinates_list[i + 1][0]},{coordinates_list[i + 1][1]}"
        start_name = location_names[i] if i < len(location_names) else "Unknown"
        end_name = location_names[i + 1] if i + 1 < len(location_names) else "Unknown"

        url = f"https://router.hereapi.com/v8/routes?transportMode=car&origin={start_coords}&destination={end_coords}&return=summary&apiKey={API_KEY}"
        response = requests.get(url).json()

        if response.get("routes"):
            summary = response["routes"][0]["sections"][0]["summary"]
            distance = summary["length"] / 1000
            duration = summary["duration"] / 60
            total_distance += distance
            total_duration += duration
            results.append(
                f"Route from {start_name} to {end_name}:\n  Distance: {distance:.2f} km\n  Duration: {duration:.2f} mins\n"
            )
        else:
            results.append(f"Route from {start_name} to {end_name} not found.\n")

    m_result_text.delete(1.0, tk.END)
    m_result_text.insert(tk.END, "\n".join(results))
    m_result_text.insert(tk.END, f"\nTotal Distance: {total_distance:.2f} km\n")
    m_result_text.insert(tk.END, f"Total Duration: {total_duration:.2f} mins\n")

def m_clear_all():
    m_map_widget.delete_all_marker()
    m_search_entry.delete(0, tk.END)
    m_result_text.delete(1.0, tk.END)
    coordinates_list.clear()
    location_names.clear()
    m_result_text.insert(tk.END, "All data cleared.\n")

def m_export_coordinates():
    if coordinates_list:
        with open("coordinates.csv", "w", newline="") as file:
            writer = csv.writer(file)
            writer.writerow(["Latitude", "Longitude"])
            writer.writerows(coordinates_list)
        m_result_text.delete(1.0, tk.END)
        m_result_text.insert(tk.END, "Coordinates exported to coordinates.csv\n")
    else:
        m_result_text.delete(1.0, tk.END)
        m_result_text.insert(tk.END, "No coordinates to export.\n")

def m_fit_bounds_manual():
    if not coordinates_list:
        return

    min_lat = min(lat for lat, lng in coordinates_list)
    max_lat = max(lat for lat, lng in coordinates_list)
    min_lng = min(lng for lat, lng in coordinates_list)
    max_lng = max(lng for lat, lng in coordinates_list)

    center_lat = (min_lat + max_lat) / 2
    center_lng = (min_lng + max_lng) / 2

    m_map_widget.set_position(center_lat, center_lng)

    lat_diff = max_lat - min_lat
    lng_diff = max_lng - min_lng
    max_diff = max(lat_diff, lng_diff)

    if max_diff < 0.05:
        zoom_level = 15
    elif max_diff < 0.1:
        zoom_level = 14
    elif max_diff < 0.5:
        zoom_level = 12
    elif max_diff < 1:
        zoom_level = 10
    elif max_diff < 5:
        zoom_level = 8
    else:
        zoom_level = 5

    m_map_widget.set_zoom(zoom_level)

# Function: Validate Email Address
def evalidate_email_address(email):
    try:
        valid = validate_email(email)
        return f"Valid email address: {valid.email}", True
    except EmailNotValidError as e:
        return f"Invalid email address: {str(e)}", False

# Function: Domain WHOIS Lookup
def edomain_whois_lookup(email):
    domain = email.split('@')[-1]
    try:
        domain_info = whois.whois(domain)
        statuses = set(status.split()[0] for status in domain_info.status)
        formatted_info = (
            f"Domain: {domain_info.domain_name}\n"
            f"Registrar: {domain_info.registrar}\n"
            f"Creation Date: {domain_info.creation_date[0]}\n"
            f"Expiration Date: {domain_info.expiration_date[0]}\n"
            f"Organization: {domain_info.org}\n"
            f"Name Servers: {', '.join(domain_info.name_servers)}\n"
            f"Status: {', '.join(statuses)}"
        )
        return f"Domain WHOIS Info:\n{formatted_info}", True
    except Exception as e:
        return f"Failed to retrieve WHOIS info: {str(e)}", False

# Function: MX Record Lookup
def emx_record_lookup(email):
    domain = email.split('@')[-1]
    try:
        answers = dns.resolver.resolve(domain, 'MX')
        mx_records = [answer.exchange.to_text() for answer in answers]
        formatted_records = "\n".join(f"{i + 1}. {record}" for i, record in enumerate(mx_records))
        return f"MX Records for {domain}:\n{formatted_records}", True
    except Exception as e:
        return f"Failed to retrieve MX records: {str(e)}", False

# Start OSINT function
def estart_osint():
    email = eemail_entry.get().strip()
    eresult_box.delete("1.0", "end")  # Clear previous results
    mail_tab.update_idletasks()

    if not email:
        eresult_box.insert("end", "Error: Email address is required.\n", "red")
        return

    eresult_box.insert("end", f"Starting OSINT for: {email}\n\n")
    estart_time = time.time()
    
    # 1. Email Validation
    evalidation_result, valid = evalidate_email_address(email)
    eresult_box.insert("end", evalidation_result + "\n\n", "green" if valid else "red")
    mail_tab.update_idletasks()
    time.sleep(0.3)

    # 2. Domain WHOIS Lookup
    ewhois_result, success = edomain_whois_lookup(email)
    eresult_box.insert("end", ewhois_result + "\n\n", "green" if success else "red")
    mail_tab.update_idletasks()
    time.sleep(0.3)

    # 3. MX Record Lookup
    emx_result, success = emx_record_lookup(email)
    eresult_box.insert("end", emx_result + "\n\n", "green" if success else "red")
    mail_tab.update_idletasks()

    eelapsed_time = time.time() - estart_time
    eresult_box.insert("end", f"Total time taken: {eelapsed_time:.2f} seconds\n", "blue")

# Threaded OSINT start function
def estart_osint_thread():
    threading.Thread(target=estart_osint).start()

# Function: Export Results
def eexport_results():
    efile_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt")])
    if efile_path:
        with open(efile_path, "w") as efile:
            efile.write(eresult_box.get("1.0", "end"))

# Clear function to reset inputs and outputs
def eclear_results():
    eemail_entry.delete(0, "end")
    eresult_box.delete("1.0", "end")


# Initialize CustomTkinter GUI
ctk.set_appearance_mode("System")  # Options: "Light", "Dark", "System"
ctk.set_default_color_theme("blue")  # Options: "blue", "dark-blue", "green"

root = ctk.CTk()
root.title("A4 Contraband")
root.geometry("1800x800")

# Notebook (Tabs)
notebook = ctk.CTkTabview(root, width=900, height=700)
notebook.pack(padx=10, pady=10, expand=True, fill="both")

network_scanner_tab = notebook.add("IP")
map_viewer_tab = notebook.add("Map")
mac_tab = notebook.add("MAC")
mail_tab = notebook.add("Mail")
number_tab = notebook.add("Phone Number")
flight_tab = notebook.add("Flight")

# Network Scanner Frame
frame = ctk.CTkFrame(network_scanner_tab, corner_radius=10)
frame.pack(side="top", padx=20, pady=20, fill="both", expand=True)

l_frame = ctk.CTkFrame(network_scanner_tab, corner_radius=10)
l_frame.pack(side="right", padx=20, pady=20, fill="both", expand=True)

ctk.CTkLabel(frame, text="IP Address:").grid(row=0, column=0, padx=10, pady=5, sticky="w")
ip_entry = ctk.CTkEntry(frame, width=200)
ip_entry.grid(row=0, column=1, padx=10, pady=5)

ctk.CTkLabel(frame, text="Start Port:").grid(row=1, column=0, padx=10, pady=5, sticky="w")
start_port_entry = ctk.CTkEntry(frame, width=100)
start_port_entry.grid(row=1, column=1, padx=10, pady=5)

ctk.CTkLabel(frame, text="End Port:").grid(row=2, column=0, padx=10, pady=5, sticky="w")
end_port_entry = ctk.CTkEntry(frame, width=100)
end_port_entry.grid(row=2, column=1, padx=10, pady=5)

# Buttons
ctk.CTkButton(frame, text="Start Scan", command=start_scan).grid(row=3, column=0, padx=10, pady=10)
ctk.CTkButton(frame, text="Stop Scan", command=stop_scan).grid(row=4, column=0, padx=10, pady=10)
ctk.CTkButton(frame, text="Identify Device", command=identify_device_button).grid(row=3, column=1, padx=10, pady=10)
ctk.CTkButton(frame, text="Clear", command=clear_all).grid(row=3, column=2, padx=10, pady=10)
ctk.CTkButton(frame, text="Start Traceroute", command=start_traceroute).grid(row=4, column=1, padx=10, pady=10)
ctk.CTkOptionMenu(frame, values=["Save as TXT", "Save as PDF", "Save as DOCX"], command=handle_dropdown).grid(row=4, column=2, padx=10, pady=10)

# Progress Bar
progress_var = IntVar()
progress_bar = ctk.CTkProgressBar(frame, variable=progress_var, mode="determinate")
progress_bar.grid(row=5, columnspan=4, pady=10, sticky="ew")

# Text Areas
text_area_frame = ctk.CTkFrame(network_scanner_tab, corner_radius=10)
text_area_frame.pack(side="bottom", padx=20, pady=10, fill="both", expand=True)

geo_text_area = ctk.CTkTextbox(text_area_frame, wrap="word", height=200, width=400, border_width=0)
geo_text_area.grid(row=0, column=0, padx=10, pady=10)
geo_text_area.insert("end", "Geolocation Information:\n")

scan_text_area = ctk.CTkTextbox(text_area_frame, wrap="word", height=200, width=400, border_width=0)
scan_text_area.grid(row=0, column=1, padx=10, pady=10)
scan_text_area.insert("end", "Port Scanning Results:\n")

vuln_text_area = ctk.CTkTextbox(text_area_frame, wrap="word", height=200, width=400, border_width=0)
vuln_text_area.grid(row=1, column=0, padx=10, pady=10)
vuln_text_area.insert("end", "Vulnerability Scanning Results:\n")

traceroute_text_area = ctk.CTkTextbox(text_area_frame, wrap="word", height=200, width=400, border_width=0)
traceroute_text_area.grid(row=1, column=1, padx=10, pady=10)
traceroute_text_area.insert("end", "Device Identification and Traceroute Results:\n")

# Create a note Label
lnote_label = ctk.CTkLabel(l_frame, text="*Note : Designed for use solely with IPv4 on a local network", font=("Arial", 18))
lnote_label.pack(side="top", pady=20)

# IP Connectons
lip_range_label = ctk.CTkLabel(l_frame, text="Enter IP Range (e.g., 192.168.1.0/24):")
lip_range_label.pack(side="top")

lip_range_entry = ctk.CTkEntry(l_frame, width=160)
lip_range_entry.pack(side="top", pady=5)

lscan_button = ctk.CTkButton(l_frame, text="Scan Network", command=lscan_network, fg_color="#1abc9c", text_color="white", border_width=0)
lscan_button.pack(side="top", pady=5)

# Results section
lactive_ips_label = ctk.CTkLabel(l_frame, text="Active IPs and Hostnames:")
lactive_ips_label.pack(side="top")

lactive_ips_text = ctk.CTkTextbox(l_frame, height=200, width=600, wrap="word", border_width=0)
lactive_ips_text.pack(side="top", pady=5)


#Frames for seperating the GUIs (mac_tab)
left_frame = ctk.CTkFrame(mac_tab)
left_frame.pack(side="left", fill="both", expand=True, padx=20, pady=20)

right_frame = ctk.CTkFrame(mac_tab)
right_frame.pack(side="right", fill="both", expand=True, padx=20, pady=20)

#Map GUI
m_control_frame = CTkFrame(map_viewer_tab, corner_radius=10)
m_control_frame.pack(side="left", fill="y", padx=10, pady=10)

m_label = CTkLabel(m_control_frame, text="Search Location:")
m_label.pack(pady=(10, 0))

m_search_entry = CTkEntry(m_control_frame, placeholder_text="Enter location")
m_search_entry.pack(pady=10)

m_search_button = CTkButton(m_control_frame, text="Search", command=m_search_location)
m_search_button.pack(pady=10)

m_reverse_geo_button = CTkButton(m_control_frame, text="Reverse Geocode", command=m_reverse_geocode)
m_reverse_geo_button.pack(pady=10)

m_add_marker_button = CTkButton(m_control_frame, text="Add Marker", command=m_add_marker)
m_add_marker_button.pack(pady=10)

m_weather_button = CTkButton(m_control_frame, text="Get Weather", command=m_get_weather)
m_weather_button.pack(pady=10)

m_calc_route_button = CTkButton(m_control_frame, text="Calculate Route", command=m_calculate_route)
m_calc_route_button.pack(pady=10)

m_fit_bounds_button = CTkButton(m_control_frame, text="Fit to Markers", command=m_fit_bounds_manual)
m_fit_bounds_button.pack(pady=10)

m_clear_button = CTkButton(m_control_frame, text="Clear All", command=m_clear_all)
m_clear_button.pack(pady=10)

m_export_button = CTkButton(m_control_frame, text="Export Coordinates", command=m_export_coordinates)
m_export_button.pack(pady=10)

m_result_label = CTkLabel(m_control_frame, text="Result:")
m_result_label.pack(pady=(20, 5))

m_result_text = tk.Text(m_control_frame, height=10, wrap=tk.WORD)
m_result_text.pack(padx=5, pady=5)

m_map_widget = TkinterMapView(map_viewer_tab, width=700, height=600, corner_radius=0)
m_map_widget.pack(side="right", fill="both", expand=True)
m_map_widget.set_position(20.5937, 78.9629)
m_map_widget.set_zoom(5)


#Mac GUI
# Create a label, entry box, and button for input
mac_label = ctk.CTkLabel(right_frame, text="Enter MAC Address:")
mac_label.pack(side="top", pady=10)

mac_entry = ctk.CTkEntry(right_frame, width=300)
mac_entry.pack(side="top", pady=10)

lookup_button = ctk.CTkButton(right_frame, text="Look Up Vendor", command=get_mac_info)
lookup_button.pack(side="top",pady=10)

# Create a text box for displaying the output
output_text = ctk.CTkTextbox(right_frame, height=200, width=600)
output_text.pack(side="top", pady=10)
output_text.configure(state="disabled")  # Disable text box initially

#ip to mac
# Create a label to display the result
result_label = ctk.CTkLabel(left_frame, text="MAC Address will appear here", font=("Arial", 12))
result_label.pack(side="bottom", pady=20)

# Create a button to trigger the MAC address lookup
search_button = ctk.CTkButton(left_frame, text="Get MAC Address", command=display_mac)
search_button.pack(side="bottom", pady=10)

# Create and place the label and entry widget
ip_entry = ctk.CTkEntry(left_frame, placeholder_text="e.g., 192.168.1.1", width=250)
ip_entry.pack(side="bottom", pady=5)

ip_label = ctk.CTkLabel(left_frame, text="Enter IP Address:", font=("Arial", 14))
ip_label.pack(side="bottom", pady=10)

# Create a note Label
note_label = ctk.CTkLabel(left_frame, text="*Note : Designed for use solely with IPv4 on a local network", font=("Arial", 18))
note_label.pack(side="top", pady=20)

# IP Connectons
ip_range_label = ctk.CTkLabel(left_frame, text="Enter IP Range (e.g., 192.168.1.0/24):")
ip_range_label.pack(side="top")

ip_range_entry = ctk.CTkEntry(left_frame, width=160)
ip_range_entry.pack(side="top", pady=5)

scan_button = ctk.CTkButton(left_frame, text="Scan Network", command=scan_network, fg_color="#1abc9c", text_color="white", border_width=0)
scan_button.pack(side="top", pady=5)

# Results section
active_ips_label = ctk.CTkLabel(left_frame, text="Active IPs and Hostnames:")
active_ips_label.pack(side="top")

active_ips_text = ctk.CTkTextbox(left_frame, height=200, width=600, wrap="word", border_width=0)
active_ips_text.pack(side="top", pady=5)

#pn GUI
# Input label and dropdown for country codes
label = ctk.CTkLabel(number_tab, text="Select Country Code:")
label.pack(pady=10)
country_code_var = ctk.StringVar(value="+91")
country_code_combobox = ctk.CTkComboBox(number_tab, variable=country_code_var, values=list(country_codes.values()))
country_code_combobox.pack(pady=5)

# Input field for phone number
entry_label = ctk.CTkLabel(number_tab, text="Enter Phone Number:")
entry_label.pack(pady=10)
entry = ctk.CTkEntry(number_tab, width=250)
entry.pack(pady=5)

# Output box
output_box = ctk.CTkTextbox(number_tab, height=250, width=600, wrap="word")
output_box.pack(pady=10)

# Buttons for actions
lookup_button = ctk.CTkButton(number_tab, text="Lookup", command=lookup_phone)
lookup_button.pack(pady=5)

clear_button = ctk.CTkButton(number_tab, text="Clear", command=clear_fields)
clear_button.pack(pady=5)

#flight GUI
# Main frame
f_main_frame = ctk.CTkFrame(flight_tab)
f_main_frame.pack(fill="both", expand=True)


# Map widget on the left
f_map_widget = TkinterMapView(f_main_frame, width=800, height=500, corner_radius=0)
f_map_widget.pack(side="left", fill="both", expand=True)
f_map_widget.set_position(20, 0)
f_map_widget.set_zoom(2)

# Right frame for controls and log
f_right_frame = ctk.CTkFrame(f_main_frame, width=200)
f_right_frame.pack(side="right", fill="y")

country_var = StringVar()

# Scrollable and searchable customtkinter combobox
f_country_menu = ctk.CTkComboBox(f_right_frame,values=countries,variable=country_var, width=300, height=35, state="normal")
f_country_menu.set("Select a Country")
f_country_menu.pack(padx=10, pady=20)

# Bind the key release event to the filtering function
f_country_menu.bind('<KeyRelease>', f_on_country_input)


f_loading_label = ctk.CTkLabel(f_right_frame, text="", font=("Helvetica", 12))
f_loading_label.pack(padx=10, pady=5)

f_view_button = ctk.CTkButton(f_right_frame, text="View Flights", command=f_view_flights)
f_view_button.pack(padx=10, pady=10)

f_log_text = ctk.CTkTextbox(f_right_frame, state="disabled")
f_log_text.pack(fill="both", expand=True, padx=10, pady=5)

# Input box for callsign search
callsign_var = StringVar()
callsign_entry = ctk.CTkEntry(f_right_frame, textvariable=callsign_var, placeholder_text="Enter Callsign")
callsign_entry.pack(padx=10, pady=5)

# Search button for callsign
f_search_button = ctk.CTkButton(f_right_frame, text="Search Callsign", command=lambda: f_search_callsign(callsign_var.get()))
f_search_button.pack(padx=10, pady=5)


# Add the Clear Button to the GUI
f_clear_button = ctk.CTkButton(f_right_frame, text="Clear All", command=f_clear_all)
f_clear_button.pack(padx=10, pady=5)

#email GUI

# Input Field for Email Address
eemail_label = ctk.CTkLabel(mail_tab, text="Enter Email Address:")
eemail_label.pack(pady=10)

eemail_entry = ctk.CTkEntry(mail_tab, width=400)
eemail_entry.pack(pady=10)

# Start OSINT Button
estart_button = ctk.CTkButton(mail_tab, text="Start OSINT", command=estart_osint_thread)
estart_button.pack(pady=10)

# Clear Button
eclear_button = ctk.CTkButton(mail_tab, text="Clear", command=eclear_results)
eclear_button.pack(pady=5)

# Export Button
eexport_button = ctk.CTkButton(mail_tab, text="Export Results", command=eexport_results)
eexport_button.pack(pady=5)

# Output Display Box
eresult_box = ctk.CTkTextbox(mail_tab, width=600, height=250)
eresult_box.pack(pady=20)

# Tag configurations for result colors
eresult_box.tag_config("green", foreground="green")
eresult_box.tag_config("red", foreground="red")
eresult_box.tag_config("blue", foreground="blue")
root.mainloop()
